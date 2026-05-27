from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)

add_heading('数据库设计文档', 1)
doc.add_paragraph('本文档基于习惯养成小程序的实际实现，包含数据库ER图说明、系统流程图说明及各数据表结构。')

add_heading('一、ER图说明', 2)

doc.add_paragraph('系统包含6个核心实体，实体间关系如下：')
doc.add_paragraph('• 用户（yonghu）1:N 学习计划（xuexijihua）：一个用户可创建多个学习计划')
doc.add_paragraph('• 学习计划（xuexijihua）1:N 计划打卡（jihuadaka）：一个计划对应多条打卡记录')
doc.add_paragraph('• 学习计划（xuexijihua）1:1 计划历史（jihualishibiao）：计划完成后归档为历史记录')
doc.add_paragraph('• 用户（yonghu）1:N 留言板（liuyanban）：一个用户可发多条留言')
doc.add_paragraph('• 管理员（admin）管理所有用户和留言数据')
doc.add_paragraph('')

add_heading('（1）用户信息实体属性图', 3)
doc.add_paragraph('用户实体包含属性：ID、创建时间、账号、密码、姓名、性别、年龄、手机、邮箱、照片、积分。用户通过账号密码登录系统，可管理学习计划和打卡。')

add_heading('（2）学习计划实体属性图', 3)
doc.add_paragraph('学习计划实体包含属性：ID、创建时间、计划标题、计划图片、开始日期、计划内容、结束时间、计划天数、完成度、账号、姓名、用户ID。通过userid与用户关联。')

add_heading('（3）计划打卡实体属性图', 3)
doc.add_paragraph('计划打卡实体包含属性：ID、创建时间、计划标题、计划图片、开始日期、计划内容、结束时间、计划天数、完成度、打卡日期、打卡天数、账号、姓名、用户ID、计划ID、补打卡。通过jihuaid与学习计划关联。')

add_heading('（4）计划历史实体属性图', 3)
doc.add_paragraph('计划历史实体包含属性：ID、创建时间、计划标题、计划图片、开始日期、计划内容、结束时间、计划天数、完成度、完成日期、账号、姓名、用户ID。记录已完成或过期的计划。')

add_heading('（5）留言板实体属性图', 3)
doc.add_paragraph('留言板实体包含属性：ID、创建时间、用户ID、用户名、留言内容、回复内容。用户发布留言，管理员可回复。')

add_heading('（6）管理员信息实体属性图', 3)
doc.add_paragraph('管理员实体包含属性：ID、用户名、密码、角色、创建时间。管理员登录后台管理用户数据和留言。')

add_heading('二、系统流程图说明', 2)

add_heading('（1）用户登录流程', 3)
doc.add_paragraph('开始 → 输入账号密码 → 判断账号是否为空（是→提示账号不能为空→返回输入）→ 判断密码是否为空（是→提示密码不能为空→返回输入）→ 检测账号密码是否正确（否→提示账号密码错误→返回输入）→ 登录成功 → 结束')

add_heading('（2）创建学习计划流程', 3)
doc.add_paragraph('开始 → 用户登录 → 点击创建计划 → 填写计划标题、内容、开始/结束日期、计划天数 → 上传计划图片（可选）→ 判断必填项是否完整（否→提示填写完整→返回）→ 提交保存 → 计划创建成功 → 结束')

add_heading('（3）计划打卡流程', 3)
doc.add_paragraph('开始 → 用户登录 → 查看我的计划列表 → 选择一个计划 → 点击打卡 → 系统记录打卡日期和天数 → 更新完成度 → 打卡成功 → 判断是否完成全部天数（是→计划标记完成→归档到历史表）→ 结束')

add_heading('（4）留言管理流程', 3)
doc.add_paragraph('开始 → 用户登录 → 进入留言板 → 输入留言内容 → 提交 → 留言保存成功 → 管理员查看留言 → 回复留言 → 结束')

add_heading('三、数据表结构', 2)

add_heading('表1 用户表（yonghu）', 3)
doc.add_paragraph('存储系统注册用户的基本信息，包括账号、密码、个人资料和积分。')
add_table(
    ['字段', '类型', '空', '默认', '注释'],
    [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
        ['zhanghao', 'varchar(200)', '否', '', '账号'],
        ['mima', 'varchar(200)', '否', '', '密码'],
        ['xingming', 'varchar(200)', '是', 'NULL', '姓名'],
        ['xingbie', 'varchar(200)', '是', 'NULL', '性别'],
        ['nianling', 'int(11)', '是', 'NULL', '年龄'],
        ['shouji', 'varchar(200)', '是', 'NULL', '手机'],
        ['youxiang', 'varchar(200)', '是', 'NULL', '邮箱'],
        ['zhaopian', 'varchar(200)', '是', 'NULL', '照片'],
        ['jifen', 'int(11)', '是', 'NULL', '积分'],
    ]
)
doc.add_paragraph('')

add_heading('表2 学习计划表（xuexijihua）', 3)
doc.add_paragraph('存储用户创建的学习计划信息，通过userid关联用户。')
add_table(
    ['字段', '类型', '空', '默认', '注释'],
    [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
        ['jihuabiaoti', 'varchar(200)', '否', '', '计划标题'],
        ['jihuatupian', 'varchar(200)', '是', 'NULL', '计划图片'],
        ['kaishiriqi', 'varchar(200)', '是', 'NULL', '开始日期'],
        ['jihuaneirong', 'longtext', '是', 'NULL', '计划内容'],
        ['jieshushijian', 'varchar(200)', '是', 'NULL', '结束时间'],
        ['jihuatianshu', 'varchar(200)', '是', 'NULL', '计划天数'],
        ['wanchengdu', 'varchar(200)', '是', 'NULL', '完成度'],
        ['zhanghao', 'varchar(200)', '是', 'NULL', '账号'],
        ['xingming', 'varchar(200)', '是', 'NULL', '姓名'],
        ['userid', 'bigint(20)', '是', 'NULL', '用户ID'],
    ]
)
doc.add_paragraph('')

add_heading('表3 计划打卡表（jihuadaka）', 3)
doc.add_paragraph('记录用户每日打卡情况，通过jihuaid关联学习计划，通过userid关联用户。')
add_table(
    ['字段', '类型', '空', '默认', '注释'],
    [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
        ['jihuabiaoti', 'varchar(200)', '否', '', '计划标题'],
        ['jihuatupian', 'varchar(200)', '是', 'NULL', '计划图片'],
        ['kaishiriqi', 'varchar(200)', '是', 'NULL', '开始日期'],
        ['jihuaneirong', 'longtext', '是', 'NULL', '计划内容'],
        ['jieshushijian', 'varchar(200)', '是', 'NULL', '结束时间'],
        ['jihuatianshu', 'varchar(200)', '是', 'NULL', '计划天数'],
        ['wanchengdu', 'varchar(200)', '是', 'NULL', '完成度'],
        ['dakariqi', 'date', '是', 'NULL', '打卡日期'],
        ['dakatianshu', 'int(11)', '是', 'NULL', '打卡天数'],
        ['zhanghao', 'varchar(200)', '是', 'NULL', '账号'],
        ['xingming', 'varchar(200)', '是', 'NULL', '姓名'],
        ['userid', 'bigint(20)', '是', 'NULL', '用户ID'],
        ['jihuaid', 'bigint(20)', '是', 'NULL', '计划ID'],
        ['budaka', 'int(11)', '是', 'NULL', '补打卡'],
    ]
)
doc.add_paragraph('')

add_heading('表4 计划历史表（jihualishibiao）', 3)
doc.add_paragraph('存储已完成或过期的学习计划归档记录，通过userid关联用户。')
add_table(
    ['字段', '类型', '空', '默认', '注释'],
    [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
        ['jihuabiaoti', 'varchar(200)', '否', '', '计划标题'],
        ['jihuatupian', 'varchar(200)', '是', 'NULL', '计划图片'],
        ['kaishiriqi', 'varchar(200)', '是', 'NULL', '开始日期'],
        ['jihuaneirong', 'longtext', '是', 'NULL', '计划内容'],
        ['jieshushijian', 'varchar(200)', '是', 'NULL', '结束时间'],
        ['jihuatianshu', 'int(11)', '是', 'NULL', '计划天数'],
        ['wanchengdu', 'varchar(200)', '是', 'NULL', '完成度'],
        ['wanchengriqi', 'varchar(200)', '是', 'NULL', '完成日期'],
        ['zhanghao', 'varchar(200)', '是', 'NULL', '账号'],
        ['xingming', 'varchar(200)', '是', 'NULL', '姓名'],
        ['userid', 'bigint(20)', '是', 'NULL', '用户ID'],
    ]
)
doc.add_paragraph('')

add_heading('表5 留言板表（liuyanban）', 3)
doc.add_paragraph('存储用户留言及管理员回复，通过userid关联用户。')
add_table(
    ['字段', '类型', '空', '默认', '注释'],
    [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
        ['userid', 'bigint(20)', '是', 'NULL', '用户ID'],
        ['username', 'varchar(200)', '是', 'NULL', '用户名'],
        ['content', 'longtext', '是', 'NULL', '留言内容'],
        ['reply', 'longtext', '是', 'NULL', '回复内容'],
    ]
)
doc.add_paragraph('')

add_heading('表6 管理员表（admin）', 3)
doc.add_paragraph('存储后台管理员的登录信息。')
add_table(
    ['字段', '类型', '空', '默认', '注释'],
    [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['username', 'varchar(200)', '否', '', '用户名'],
        ['password', 'varchar(200)', '否', '', '密码'],
        ['role', 'varchar(200)', '是', 'NULL', '角色'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
    ]
)

output_path = r'F:\26dan\习惯小程序\数据库设计文档.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
