import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

matplotlib.rcParams['font.family'] = ['SimSun']
matplotlib.rcParams['axes.unicode_minus'] = False

OUT_DIR = r'F:\26dan\习惯小程序\doc_images'
os.makedirs(OUT_DIR, exist_ok=True)

def draw_er(entity_name, attributes, filename):
    n = len(attributes)

    rw, rh = 1.0, 0.4
    ew, eh = 0.75, 0.3

    min_gap = 2 * ew + 0.3
    radius = max(2.8, min_gap * n / (2 * np.pi))

    bound = radius + ew + 0.5
    fig, ax = plt.subplots(1, 1, figsize=(bound * 2, bound * 2))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.set_xlim(-bound, bound)
    ax.set_ylim(-bound, bound)
    ax.set_aspect('equal')
    ax.axis('off')

    rect = mpatches.Rectangle((-rw, -rh), 2*rw, 2*rh,
                               facecolor='white', edgecolor='black', linewidth=2.5)
    ax.add_patch(rect)
    ax.text(0, 0, entity_name, ha='center', va='center',
            fontsize=24, fontweight='bold', fontfamily='SimSun')

    angles = np.linspace(0, 2 * np.pi, n, endpoint=False) - np.pi / 2

    for i, attr in enumerate(attributes):
        ang = angles[i]
        cos_a, sin_a = np.cos(ang), np.sin(ang)
        cx = radius * cos_a
        cy = radius * sin_a

        ellipse = mpatches.Ellipse((cx, cy), 2*ew, 2*eh,
                                    facecolor='white', edgecolor='black', linewidth=2)
        ax.add_patch(ellipse)
        ax.text(cx, cy, attr, ha='center', va='center',
                fontsize=18, fontfamily='SimSun')

        t_candidates = []
        if abs(cos_a) > 1e-10:
            t_candidates.append(rw / abs(cos_a))
        if abs(sin_a) > 1e-10:
            t_candidates.append(rh / abs(sin_a))
        t_rect = min(t_candidates) if t_candidates else 1.0
        sx = t_rect * cos_a
        sy = t_rect * sin_a

        dx_e, dy_e = -cx, -cy
        denom = np.sqrt(dx_e**2 / ew**2 + dy_e**2 / eh**2)
        if denom > 1e-10:
            ex = cx + dx_e / denom
            ey = cy + dy_e / denom
        else:
            ex, ey = cx, cy

        ax.plot([sx, ex], [sy, ey], 'k-', linewidth=1.5)

    plt.tight_layout(pad=0)
    path = os.path.join(OUT_DIR, filename)
    plt.savefig(path, dpi=250, bbox_inches='tight',
                facecolor='white', edgecolor='none', pad_inches=0.2)
    plt.close()
    return path

def draw_flowchart(steps, filename):
    """
    steps 格式:
    - ('start', text)
    - ('end', text)
    - ('process', text)
    - ('decision', text, side_label, side_dir, side_text, return_idx, main_label)
      side_dir: 'left'/'right', return_idx: 回到第几步(>=0)或-1(汇合到下一步)
    """
    n = len(steps)
    step_h = 2.2
    side_x_offset = 3.8

    proc_hw, proc_hh = 1.5, 0.4
    start_hw, start_hh = 1.2, 0.35
    dec_hw, dec_hh = 1.6, 0.6
    side_hw, side_hh = 1.3, 0.4

    total_h = n * step_h + 2
    fig, ax = plt.subplots(1, 1, figsize=(8, total_h * 0.6))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.set_xlim(-6.5, 6.5)
    ax.set_ylim(-total_h + 0.5, 1.5)
    ax.axis('off')

    positions = []
    y = 0

    for i, step in enumerate(steps):
        stype = step[0]
        text = step[1]
        cx = 0
        positions.append((cx, y))

        if stype in ('start', 'end'):
            shape = FancyBboxPatch((cx - start_hw, y - start_hh), 2*start_hw, 2*start_hh,
                                   boxstyle="round,pad=0.15",
                                   facecolor='white', edgecolor='black', linewidth=2)
            ax.add_patch(shape)
            ax.text(cx, y, text, ha='center', va='center', fontsize=16, fontfamily='SimSun')
            s_top, s_bot = y + start_hh, y - start_hh

        elif stype == 'process':
            rect = mpatches.Rectangle((cx - proc_hw, y - proc_hh), 2*proc_hw, 2*proc_hh,
                                       facecolor='white', edgecolor='black', linewidth=2)
            ax.add_patch(rect)
            ax.text(cx, y, text, ha='center', va='center', fontsize=14, fontfamily='SimSun')
            s_top, s_bot = y + proc_hh, y - proc_hh

        elif stype == 'decision':
            diamond = plt.Polygon([(cx, y+dec_hh), (cx+dec_hw, y), (cx, y-dec_hh), (cx-dec_hw, y)],
                                   facecolor='white', edgecolor='black', linewidth=2, closed=True)
            ax.add_patch(diamond)
            ax.text(cx, y, text, ha='center', va='center', fontsize=12, fontfamily='SimSun')
            s_top, s_bot = y + dec_hh, y - dec_hh

            if len(step) > 2:
                sl = step[2]
                sd = step[3]
                st = step[4]
                sr = step[5]
                ml = step[6]
                dsign = -1 if sd == 'left' else 1
                sx = side_x_offset * dsign

                side_rect = mpatches.Rectangle((sx - side_hw, y - side_hh), 2*side_hw, 2*side_hh,
                                                facecolor='white', edgecolor='black', linewidth=1.5)
                ax.add_patch(side_rect)
                ax.text(sx, y, st, ha='center', va='center', fontsize=11, fontfamily='SimSun')

                d_edge = cx + dec_hw * dsign
                s_edge = sx - side_hw * dsign
                ax.annotate('', xy=(s_edge, y), xytext=(d_edge, y),
                           arrowprops=dict(arrowstyle='->', color='black', lw=1.5))
                lx = (d_edge + s_edge) / 2
                ax.text(lx, y + 0.15, sl, ha='center', va='bottom', fontsize=13, fontfamily='SimSun')

                if sr >= 0:
                    ty = positions[sr][1]
                    s_out = sx + side_hw * dsign
                    loop_x = s_out + 0.6 * dsign
                    t_edge = proc_hw * dsign
                    ax.plot([s_out, loop_x], [y, y], 'k-', linewidth=1.5)
                    ax.plot([loop_x, loop_x], [y, ty], 'k-', linewidth=1.5)
                    ax.annotate('', xy=(t_edge, ty), xytext=(loop_x, ty),
                               arrowprops=dict(arrowstyle='->', color='black', lw=1.5))
                elif sr == -1:
                    next_y = y - step_h
                    s_out = sx + side_hw * dsign
                    merge_y = next_y + proc_hh + 0.3
                    ax.plot([s_out, s_out], [y - side_hh, merge_y], 'k-', linewidth=1.5)
                    ax.annotate('', xy=(cx, merge_y), xytext=(s_out, merge_y),
                               arrowprops=dict(arrowstyle='->', color='black', lw=1.5))

                ax.text(cx + 0.12, s_bot - 0.12, ml, ha='left', va='top',
                        fontsize=13, fontfamily='SimSun')

        if i > 0:
            pt = steps[i-1][0]
            py = positions[i-1][1]
            if pt in ('start', 'end'):
                fy = py - start_hh
            elif pt == 'process':
                fy = py - proc_hh
            elif pt == 'decision':
                fy = py - dec_hh
            ax.annotate('', xy=(cx, s_top), xytext=(cx, fy),
                       arrowprops=dict(arrowstyle='->', color='black', lw=1.5))

        y -= step_h

    plt.tight_layout(pad=0)
    path = os.path.join(OUT_DIR, filename)
    plt.savefig(path, dpi=250, bbox_inches='tight',
                facecolor='white', edgecolor='none', pad_inches=0.2)
    plt.close()
    return path

er_configs = [
    ('用户', ['ID', '创建时间', '账号', '密码', '姓名', '性别', '年龄', '手机', '邮箱', '照片', '积分'], 'er_yonghu.png'),
    ('学习计划', ['ID', '创建时间', '计划标题', '计划图片', '开始日期', '计划内容', '结束时间', '计划天数', '完成度', '账号', '姓名'], 'er_xuexijihua.png'),
    ('计划打卡', ['ID', '创建时间', '计划标题', '打卡日期', '打卡天数', '计划天数', '完成度', '补打卡', '计划ID', '用户ID'], 'er_jihuadaka.png'),
    ('计划历史', ['ID', '创建时间', '计划标题', '计划图片', '开始日期', '计划内容', '结束时间', '计划天数', '完成度', '完成日期', '用户ID'], 'er_jihualishibiao.png'),
    ('留言板', ['ID', '创建时间', '用户ID', '用户名', '留言内容', '回复内容'], 'er_liuyanban.png'),
    ('管理员', ['ID', '用户名', '密码', '角色', '创建时间'], 'er_admin.png'),
]

print('生成 ER 图...')
er_paths = []
for name, attrs, fname in er_configs:
    p = draw_er(name, attrs, fname)
    er_paths.append(p)
    print(f'  {name} -> {fname}')

print('生成流程图...')
login_flow = [
    ('start', '开始'),
    ('process', '输入用户名密码'),
    ('decision', '判断用户名\n是否为空', '是', 'left', '提示用户名\n不能为空', 1, '否'),
    ('decision', '判断密码\n是否为空', '是', 'right', '提示密码\n不能为空', 1, '否'),
    ('decision', '检测用户名密码\n是否正确', '否', 'right', '提示用户名\n密码错误', 1, '是'),
    ('process', '登录成功'),
    ('end', '结束'),
]
flow1 = draw_flowchart(login_flow, 'flow_login.png')

plan_flow = [
    ('start', '开始'),
    ('process', '用户登录'),
    ('process', '点击创建计划'),
    ('process', '填写计划信息'),
    ('decision', '必填项\n是否完整', '否', 'right', '提示填写\n必填项', 3, '是'),
    ('process', '提交保存'),
    ('process', '计划创建成功'),
    ('end', '结束'),
]
flow2 = draw_flowchart(plan_flow, 'flow_plan.png')

daka_flow = [
    ('start', '开始'),
    ('process', '用户登录'),
    ('process', '查看计划列表'),
    ('process', '选择计划'),
    ('process', '点击打卡'),
    ('process', '记录打卡\n更新完成度'),
    ('decision', '是否完成\n全部天数', '是', 'right', '归档到\n历史表', -1, '否'),
    ('process', '打卡完成'),
    ('end', '结束'),
]
flow3 = draw_flowchart(daka_flow, 'flow_daka.png')

liuyan_flow = [
    ('start', '开始'),
    ('process', '用户登录'),
    ('process', '进入留言板'),
    ('process', '输入留言内容'),
    ('decision', '留言内容\n是否为空', '是', 'left', '提示输入\n留言内容', 3, '否'),
    ('process', '提交留言'),
    ('process', '管理员查看\n并回复'),
    ('end', '结束'),
]
flow4 = draw_flowchart(liuyan_flow, 'flow_liuyan.png')

print('生成 docx...')
doc = Document()
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '宋体'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def add_centered_image(path, width=Inches(5.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)

def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(10)
        run.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading('数据库设计文档', 1)
p = doc.add_paragraph('本文档基于习惯养成小程序的实际实现，包含数据库ER图、系统流程图及各数据表结构。')
for run in p.runs:
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading('一、实体属性图（ER图）', 2)

er_titles = [
    '图4.1 用户信息实体属性图',
    '图4.2 学习计划实体属性图',
    '图4.3 计划打卡实体属性图',
    '图4.4 计划历史实体属性图',
    '图4.5 留言板实体属性图',
    '图4.6 管理员信息实体属性图',
]
er_descs = [
    '用户实体存储注册用户的基本信息，包括账号密码、个人资料和积分。',
    '学习计划实体记录用户创建的计划信息，通过用户ID与用户关联。',
    '计划打卡实体记录每日打卡情况，通过计划ID关联学习计划。',
    '计划历史实体存储已完成或过期的计划归档记录。',
    '留言板实体存储用户留言及管理员回复。',
    '管理员实体存储后台管理员的登录信息。',
]

for i, (path, title, desc) in enumerate(zip(er_paths, er_titles, er_descs)):
    p = doc.add_paragraph(f'（{i+1}）{desc}')
    for run in p.runs:
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    add_centered_image(path, width=Inches(5.5))
    add_caption(title)
    doc.add_paragraph('')

add_heading('二、系统流程图', 2)

flow_data = [
    (flow1, '图5.1 用户登录流程图', '用户输入账号密码后，系统依次验证账号和密码是否为空、是否正确，验证通过后登录成功。'),
    (flow2, '图5.2 创建学习计划流程图', '用户填写计划标题、内容、日期等信息，系统验证必填项完整性，通过后保存计划。'),
    (flow3, '图5.3 计划打卡流程图', '用户选择计划并打卡，系统更新完成度，完成全部天数后自动归档到历史表。'),
    (flow4, '图5.4 留言管理流程图', '用户提交留言后，管理员在后台查看并进行回复。'),
]

for i, (path, title, desc) in enumerate(flow_data):
    p = doc.add_paragraph(f'（{i+1}）{desc}')
    for run in p.runs:
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    add_centered_image(path, width=Inches(4.5))
    add_caption(title)
    doc.add_paragraph('')

add_heading('三、数据表结构', 2)

tables_data = [
    ('表4.1 用户表（yonghu）', '存储系统注册用户的基本信息。', [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
        ['zhanghao', 'varchar(200)', '否', '', '账号'],
        ['mima', 'varchar(200)', '否', '', '密码'],
        ['xingming', 'varchar(200)', '是', 'NULL', '姓名'],
        ['xingbie', 'varchar(200)', '是', 'NULL', '性别'],
        ['nianling', 'int(11)', '是', 'NULL', '年龄'],
        ['shouji', 'varchar(200)', '是', 'NULL', '手机'],
        ['youxiang', 'varchar(200)', '是', 'NULL', '邮箱'],
        ['zhaopian', 'varchar(200)', '是', 'NULL', '照片'],
        ['jifen', 'int(11)', '是', 'NULL', '积分'],
    ]),
    ('表4.2 学习计划表（xuexijihua）', '存储用户创建的学习计划信息。', [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
        ['jihuabiaoti', 'varchar(200)', '否', '', '计划标题'],
        ['jihuatupian', 'varchar(200)', '是', 'NULL', '计划图片'],
        ['kaishiriqi', 'varchar(200)', '是', 'NULL', '开始日期'],
        ['jihuaneirong', 'longtext', '是', 'NULL', '计划内容'],
        ['jieshushijian', 'varchar(200)', '是', 'NULL', '结束时间'],
        ['jihuatianshu', 'varchar(200)', '是', 'NULL', '计划天数'],
        ['wanchengdu', 'varchar(200)', '是', 'NULL', '完成度'],
        ['zhanghao', 'varchar(200)', '是', 'NULL', '账号'],
        ['xingming', 'varchar(200)', '是', 'NULL', '姓名'],
        ['userid', 'bigint(20)', '是', 'NULL', '用户ID'],
    ]),
    ('表4.3 计划打卡表（jihuadaka）', '记录用户每日打卡情况。', [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
        ['jihuabiaoti', 'varchar(200)', '否', '', '计划标题'],
        ['jihuatupian', 'varchar(200)', '是', 'NULL', '计划图片'],
        ['kaishiriqi', 'varchar(200)', '是', 'NULL', '开始日期'],
        ['jihuaneirong', 'longtext', '是', 'NULL', '计划内容'],
        ['jieshushijian', 'varchar(200)', '是', 'NULL', '结束时间'],
        ['jihuatianshu', 'varchar(200)', '是', 'NULL', '计划天数'],
        ['wanchengdu', 'varchar(200)', '是', 'NULL', '完成度'],
        ['dakariqi', 'date', '是', 'NULL', '打卡日期'],
        ['dakatianshu', 'int(11)', '是', 'NULL', '打卡天数'],
        ['zhanghao', 'varchar(200)', '是', 'NULL', '账号'],
        ['xingming', 'varchar(200)', '是', 'NULL', '姓名'],
        ['userid', 'bigint(20)', '是', 'NULL', '用户ID'],
        ['jihuaid', 'bigint(20)', '是', 'NULL', '计划ID'],
        ['budaka', 'int(11)', '是', 'NULL', '补打卡'],
    ]),
    ('表4.4 计划历史表（jihualishibiao）', '存储已完成或过期的学习计划归档记录。', [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
        ['jihuabiaoti', 'varchar(200)', '否', '', '计划标题'],
        ['jihuatupian', 'varchar(200)', '是', 'NULL', '计划图片'],
        ['kaishiriqi', 'varchar(200)', '是', 'NULL', '开始日期'],
        ['jihuaneirong', 'longtext', '是', 'NULL', '计划内容'],
        ['jieshushijian', 'varchar(200)', '是', 'NULL', '结束时间'],
        ['jihuatianshu', 'int(11)', '是', 'NULL', '计划天数'],
        ['wanchengdu', 'varchar(200)', '是', 'NULL', '完成度'],
        ['wanchengriqi', 'varchar(200)', '是', 'NULL', '完成日期'],
        ['zhanghao', 'varchar(200)', '是', 'NULL', '账号'],
        ['xingming', 'varchar(200)', '是', 'NULL', '姓名'],
        ['userid', 'bigint(20)', '是', 'NULL', '用户ID'],
    ]),
    ('表4.5 留言板表（liuyanban）', '存储用户留言及管理员回复。', [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
        ['userid', 'bigint(20)', '是', 'NULL', '用户ID'],
        ['username', 'varchar(200)', '是', 'NULL', '用户名'],
        ['content', 'longtext', '是', 'NULL', '留言内容'],
        ['reply', 'longtext', '是', 'NULL', '回复内容'],
    ]),
    ('表4.6 管理员表（admin）', '存储后台管理员的登录信息。', [
        ['id (主键)', 'bigint(20)', '否', '', '主键'],
        ['username', 'varchar(200)', '否', '', '用户名'],
        ['password', 'varchar(200)', '否', '', '密码'],
        ['role', 'varchar(200)', '是', 'NULL', '角色'],
        ['addtime', 'timestamp', '否', 'CURRENT_TIMESTAMP', '创建时间'],
    ]),
]

for title, desc, rows in tables_data:
    add_heading(title, 3)
    p = doc.add_paragraph(desc)
    for run in p.runs:
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    add_table(['字段', '类型', '空', '默认', '注释'], rows)
    doc.add_paragraph('')

output_path = r'F:\26dan\习惯小程序\数据库设计文档_v3.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
